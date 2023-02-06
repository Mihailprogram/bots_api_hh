from docx import Document
from docx.shared import Mm
import os
def word():
    directory = "C:/Users/Химачи/Desktop/Курс/"
    files = os.listdir(directory)
    mages = filter(lambda x: x.endswith('.jpg'), files)
    mas=[]
    for i in mages:
            mas.append(i)
    dock=Document()
    dock.paragraphs[:5]
    for i in mas:
        dock.add_picture(i, width = Mm(150))

    dock.save("BDZ.docx")
